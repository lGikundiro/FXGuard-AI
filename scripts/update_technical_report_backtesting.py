"""Adapt the Word technical report to the current multicurrency backtesting results."""
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
REPORT_PATH = ROOT / "FXGuard_AI_Technical_Report_converted.docx"
BACKUP_PATH = ROOT / "FXGuard_AI_Technical_Report_before_multicurrency_backtesting.docx"
TEMP_PATH = ROOT / "FXGuard_AI_Technical_Report_converted.updating.docx"


def replace_paragraph_text(paragraph, text: str) -> None:
    """Replace text while retaining the direct formatting of the first run."""
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def replace_cell_text(cell, text: str) -> None:
    paragraph = cell.paragraphs[0]
    replace_paragraph_text(paragraph, text)
    for extra_paragraph in cell.paragraphs[1:]:
        extra_paragraph._element.getparent().remove(extra_paragraph._element)


def set_table_row(row, values) -> None:
    if len(row.cells) != len(values):
        raise ValueError(f"Expected {len(row.cells)} values, received {len(values)}.")
    for cell, value in zip(row.cells, values):
        replace_cell_text(cell, str(value))


def insert_paragraphs_after(document, anchor_element, items) -> None:
    current_anchor = anchor_element
    for style, text in items:
        paragraph = document.add_paragraph(text, style=style)
        current_anchor.addnext(paragraph._p)
        current_anchor = paragraph._p


def update_report() -> None:
    if not REPORT_PATH.exists():
        raise FileNotFoundError(REPORT_PATH)
    if not BACKUP_PATH.exists():
        shutil.copy2(REPORT_PATH, BACKUP_PATH)

    document = Document(REPORT_PATH)
    if any(p.text.strip() == "Automated Software Testing" for p in document.paragraphs):
        current_updates = {
            5: (
                "Rwanda-based importers that earn RWF and owe foreign-currency supplier "
                "payments face uncertain RWF costs when rates move before settlement. "
                "FXGuard AI investigates whether short-term risk classification can add "
                "payment-planning support. It does not stop depreciation, hedge exposure, "
                "or resolve the underlying financial risk."
            ),
            6: (
                "FXGuard AI is an experimental web-based decision-support prototype. It "
                "uses official USD/RWF, EUR/RWF, and KES/RWF histories and the user's "
                "invoice payment date to classify 1-to-100-day depreciation pressure as "
                "Low, Medium, or High."
            ),
            9: (
                "The model dataset uses actual BNR posting rows only. A forward-filled "
                "daily calendar remains available for charts, but weekends and other "
                "non-posting dates are excluded from model training."
            ),
            12: (
                "Every integer horizon from 1 to 100 days uses the first official BNR "
                "posting on or after its calendar target date. Within evaluation, Low, "
                "Medium, and High thresholds are derived separately for every horizon "
                "from the earlier training window only."
            ),
            21: (
                "Logistic Regression, Random Forest, and XGBoost were compared "
                "independently for USD, EUR, and KES across the full 1-to-100-day "
                "horizon surface. Model selection was based on time-aware historical "
                "backtesting rather than the final holdout."
            ),
            24: (
                "Three expanding rolling-origin folds were used for every currency and "
                "the full 1-to-100-day horizon surface. Training rows whose future outcome "
                "date reached the following test period were purged, creating up to a "
                "100-day boundary gap."
            ),
            25: (
                "*The final USD flexible-horizon holdout contains only Low-risk labels. "
                "Its 100% accuracy is not evidence that the model distinguishes all three classes."
            ),
            27: (
                "All 42 automated software tests passed. They cover the application, data "
                "validation, 1-to-100-day feature construction, payment-date limits, model "
                "outputs, reliability gate, forward-only evaluation, API, frontend, and "
                "account security foundations."
            ),
            29: (
                "None of the three flexible classifiers passed the declared reliability gate. The "
                "software is a functioning experimental prototype, but the current "
                "evidence does not support operational financial reliance."
            ),
            32: (
                "Logistic Regression was selected for USD, XGBoost for EUR, and Random "
                "Forest for KES. Selection used mean rolling balanced accuracy, then "
                "Macro F1 and accuracy."
            ),
            33: (
                "Pooled rolling balanced accuracy was 0.4472 for USD, 0.2987 for EUR, and "
                "0.4016 for KES. Macro F1 was also below the declared gate for all three."
            ),
            35: (
                "More complex algorithms did not reliably overcome the limited predictive "
                "signal in rate-history features. Algorithm choice alone is therefore not "
                "the main solution to the reliability problem."
            ),
            36: (
                "The pre-declared project gate requires mean balanced accuracy of at least "
                "0.55, mean Macro F1 of at least 0.45, and at least 0.05 improvement over "
                "the baseline. All current models remain experimental."
            ),
            37: "Uncalibrated Model Score",
            38: (
                "The classifier returns relative scores for Low, Medium, and High. The "
                "largest score accompanies the selected class, but it has not been shown "
                "to equal the real-world frequency of that outcome."
            ),
            39: (
                "The interface therefore labels this value an uncalibrated model score, "
                "not a likelihood probability or confidence guarantee."
            ),
            43: (
                "The user selects the expected invoice payment date. The browser accepts "
                "tomorrow through 100 days ahead, and the API independently enforces the "
                "same limit using the current Rwanda date."
            ),
            45: (
                "The backend calculates the number of calendar days until payment, loads "
                "the selected currency's flexible model, and supplies horizon_days with "
                "the latest official-posting features."
            ),
            46: (
                "The model returns a risk class and uncalibrated class scores. The "
                "interface prefixes the leading score with an approximation sign, such "
                "as ≈ 72%, and shows that the model is experimental."
            ),
        }
        for index, text in current_updates.items():
            replace_paragraph_text(document.paragraphs[index], text)

        current_results = [
            ("Currency", "Period", "Selected model", "Backtest accuracy", "Balanced accuracy", "Macro F1", "Holdout accuracy"),
            ("USD/RWF", "1–100 days", "Logistic Regression", "0.6386", "0.4472", "0.2801", "1.0000*"),
            ("EUR/RWF", "1–100 days", "XGBoost", "0.3084", "0.2987", "0.1680", "0.8164"),
            ("KES/RWF", "1–100 days", "Random Forest", "0.4982", "0.4016", "0.2470", "0.2481"),
        ]
        for row, values in zip(document.tables[6].rows, current_results):
            set_table_row(row, values)
        for row_index in range(len(document.tables[6].rows) - 1, len(current_results) - 1, -1):
            row = document.tables[6].rows[row_index]
            row._element.getparent().remove(row._element)

        replace_cell_text(
            document.tables[2].rows[6].cells[1],
            "Forward-filled dates are retained for charts only; model training uses actual BNR postings.",
        )
        replace_cell_text(
            document.tables[2].rows[5].cells[1],
            "Official-posting observations, engineered rate-history features, horizon_days, and horizon-specific 1-to-100-day labels.",
        )
        replace_cell_text(
            document.tables[1].rows[5].cells[1],
            "Experimental risk class, uncalibrated model scores, RWF cost scenario, drivers, and planning considerations.",
        )
        replace_cell_text(
            document.tables[1].rows[4].cells[1],
            "USD, EUR, and KES against RWF; user-selected payment dates 1–100 days ahead; FastAPI backend; web frontend.",
        )
        replace_cell_text(
            document.tables[5].rows[6].cells[1],
            "Receives currency, amount, and payment_date; validates 1–100 days; returns experimental risk classification, approximate uncalibrated scores, reliability status, cost estimates, and planning considerations.",
        )
        replace_cell_text(
            document.tables[5].rows[3].cells[1],
            "Loads one horizon-aware model for the selected currency and predicts from the latest features plus horizon_days.",
        )
        score_rows = {
            0: ("Item", "Description"),
            1: ("Model-score output", "Example display: ≈ 92% for the selected class."),
            2: ("Selected class", "The class with the largest model score."),
            3: ("Top model score", "An uncalibrated relative score, not a verified outcome probability."),
            4: ("Why it matters", "It exposes how strongly the model favors one class while retaining an explicit reliability warning."),
        }
        for row_index, values in score_rows.items():
            set_table_row(document.tables[7].rows[row_index], values)
        assumption_updates = {
            3: ("Preprocessing decision", "Model training excludes forward-filled non-posting dates; the daily calendar is display-only."),
            5: ("Limitation 2", "No current classifier passes the project-defined reliability gate; final USD holdouts also contain only Low risk."),
            7: ("Future work", "Test publication-lagged macro features, collect new outcomes, calibrate scores, and test importer comprehension directly."),
        }
        for row_index, (label, description) in assumption_updates.items():
            row = document.tables[8].rows[row_index]
            replace_cell_text(row.cells[0], label)
            replace_cell_text(row.cells[1], label)
            replace_cell_text(row.cells[2], description)
            replace_cell_text(row.cells[3], description)
        replace_cell_text(
            document.tables[4].rows[9].cells[1],
            "Number of recent days when the selected foreign-currency/RWF rate increased",
        )
        if not any(p.text.strip() == "Defense Panel Response Addendum" for p in document.paragraphs):
            document.add_heading("Defense Panel Response Addendum", level=1)
            document.add_paragraph(
                "The revised problem statement does not claim that importers universally "
                "cannot interpret exchange-rate data. Available Rwanda-specific literature "
                "supports foreign-exchange exposure and constrained firm capacity, but not "
                "that stronger importer-specific interpretation claim. Importer comprehension "
                "will be tested empirically in future user research."
            )
            document.add_paragraph(
                "Future macro-feature experiments will use inflation, policy, fuel/commodity, "
                "and trade variables only after their real publication dates. Low-frequency "
                "values may be carried forward with an age indicator, but never backfilled "
                "into dates when they were not yet public."
            )
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs[1:]:
                        if not paragraph.text.strip():
                            paragraph._element.getparent().remove(paragraph._element)
        document.save(TEMP_PATH)
        Document(TEMP_PATH)
        TEMP_PATH.replace(REPORT_PATH)
        print("Technical report updated with panel-response evaluation and normalized spacing.")
        return
    if len(document.tables) < 9 or len(document.paragraphs) < 54:
        raise ValueError("The technical report structure is not the expected version.")

    paragraphs = document.paragraphs
    tables = document.tables

    paragraph_updates = {
        1: "Multicurrency exchange-rate risk forecasting and decision-support prototype",
        5: (
            "Many importers in Rwanda earn revenue in Rwandan francs but pay international "
            "suppliers in foreign currencies. When the Rwandan franc weakens before a "
            "supplier payment date, the same invoice becomes more expensive in RWF. "
            "FXGuard AI addresses this problem by converting historical exchange-rate "
            "behaviour into an interpretable risk classification."
        ),
        6: (
            "FXGuard AI is a web-based decision-support prototype for Rwanda-based "
            "importers. It uses official USD/RWF, EUR/RWF, and KES/RWF exchange-rate "
            "history, engineered time-series features, and separate classifiers to "
            "estimate Low, Medium, or High depreciation risk over a 7-day or 14-day horizon."
        ),
        9: (
            "The current multicurrency dataset was prepared from official National Bank "
            "of Rwanda exchange-rate exports for USD, EUR, and KES. Each currency has a "
            "continuous daily calendar, engineered features, future risk labels, and "
            "separate 7-day and 14-day model-ready datasets."
        ),
        12: (
            "Labels are created by looking ahead within each currency's historical series. "
            "A 7-day label compares the current mid-rate with the rate seven days later; "
            "a 14-day label uses fourteen days later. The resulting RWF depreciation is "
            "converted into Low, Medium, or High risk using currency-specific thresholds."
        ),
        21: (
            "Logistic Regression, Random Forest, and XGBoost were compared independently "
            "for USD, EUR, and KES over both 7-day and 14-day horizons. Model selection "
            "was based on time-aware historical backtesting rather than the final holdout."
        ),
        22: (
            "Accuracy reports the overall share classified correctly. Balanced Accuracy "
            "gives equal importance to Low, Medium, and High risk, while Macro F1 combines "
            "precision and recall with equal weight for each class. Balanced Accuracy was "
            "the primary model-selection measure because risk classes were not evenly "
            "distributed in every historical period."
        ),
        23: "Multicurrency Backtesting Method",
        24: (
            "Three expanding-window rolling-origin folds were used for every currency and "
            "horizon. In each fold, the model learned from an earlier period and was tested "
            "on the following 10% of observations. All folds stayed within the first 80% "
            "of the timeline, while the final 20% remained excluded from model selection. "
            "A 7-row or 14-row purge gap separated training and testing so labels derived "
            "from future rates could not cross the boundary. After selection and holdout "
            "evaluation, the winning production model was refitted on all labeled history."
        ),
        26: "Model Selection Outcome",
        27: (
            "Logistic Regression was selected for five of the six currency/horizon tasks. "
            "Random Forest was selected for EUR/RWF over 7 days. Selection used mean rolling "
            "balanced accuracy, followed by Macro F1 and accuracy; final holdout performance "
            "did not influence the choice."
        ),
        28: (
            "USD/RWF produced the strongest historical result, with 7-day balanced accuracy "
            "of 0.5173. This indicates some useful historical signal, but not enough to treat "
            "individual predictions as certain."
        ),
        29: (
            "EUR/RWF and KES/RWF were generally close to the one-third balanced-accuracy "
            "reference point for a three-class task, showing limited separation between "
            "Low, Medium, and High risk with the current features."
        ),
        30: (
            "XGBoost was not selected for any task. Greater algorithm complexity therefore "
            "did not overcome the main challenge of finding stable short-term patterns in "
            "the available exchange-rate history."
        ),
        31: (
            "The results support an experimental decision-support prototype rather than a "
            "guaranteed forecasting system. Predictions should be considered alongside the "
            "user's business context and other financial information."
        ),
        34: (
            "For example, if a selected foreign-currency payment receives Low risk with 92% "
            "model confidence, the application displays Low and the full probability "
            "distribution. This is more transparent than displaying only a single label, "
            "but the score is not a guaranteed real-world probability."
        ),
        36: (
            "The dashboard shows the latest official rate and data freshness for the "
            "selected USD, EUR, or KES currency pair."
        ),
        37: (
            "The user opens Payment Check, selects the supplier-payment currency, and enters "
            "the invoice amount."
        ),
        42: (
            "The application calculates the estimated RWF cost and a possible additional "
            "cost for the selected currency."
        ),
    }
    for index, text in paragraph_updates.items():
        replace_paragraph_text(paragraphs[index], text)

    overview_rows = [
        ("Detail", "Description"),
        (
            "Problem",
            "Short-term USD/RWF, EUR/RWF, and KES/RWF movements can increase import costs, reduce margins, and create payment-planning uncertainty.",
        ),
        (
            "Goal",
            "Classify short-term RWF depreciation risk as Low, Medium, or High for each supported supplier-payment currency.",
        ),
        (
            "Users",
            "Rwanda-based importers, small business owners, accountants, and finance professionals.",
        ),
        (
            "Current scope",
            "USD, EUR, and KES against RWF; official BNR history; 7-day and 14-day horizons; FastAPI backend; web frontend.",
        ),
        (
            "Output",
            "Risk class, probability distribution, model-confidence score, current RWF cost, possible extra cost, key drivers, and recommendations.",
        ),
    ]
    for row, values in zip(tables[1].rows, overview_rows):
        set_table_row(row, values)

    data_rows = [
        ("Item", "Description"),
        ("Source", "Official National Bank of Rwanda Excel exchange-rate exports."),
        ("Currency pairs", "USD/RWF, EUR/RWF, and KES/RWF."),
        (
            "Period",
            "Official observations from 4 January 2022 to 17 July 2026; model-ready dates vary slightly by forecast horizon.",
        ),
        (
            "Raw fields",
            "date, currency, buying_rate, average_rate, selling_rate, source, and rate type.",
        ),
        (
            "Prepared fields",
            "Daily calendar, mid-rate, engineered features, and currency-specific 7-day and 14-day labels.",
        ),
        (
            "Daily calendar approach",
            "Missing non-posting days were forward-filled within each currency's own history; no values were created before its first official observation.",
        ),
    ]
    for row, values in zip(tables[2].rows, data_rows):
        set_table_row(row, values)

    replace_cell_text(
        tables[4].rows[9].cells[1],
        "Number of recent days when the selected foreign-currency/RWF rate increased",
    )
    replace_cell_text(
        tables[4].rows[9].cells[2],
        "Shows how persistent recent RWF weakening pressure has been for the selected currency.",
    )

    function_rows = [
        ("Function / endpoint", "Role in the prototype"),
        (
            "load_daily_calendar()",
            "Loads the prepared daily exchange-rate series for the selected currency.",
        ),
        (
            "load_features()",
            "Loads the engineered multicurrency model-ready feature dataset.",
        ),
        (
            "model_predict()",
            "Loads the currency- and horizon-specific model and returns a risk class and probabilities.",
        ),
        (
            "/api/latest-rate",
            "Returns the latest official exchange rate for USD, EUR, or KES against RWF.",
        ),
        (
            "/api/data-freshness",
            "Reports the latest imported official date and data age for a selected currency.",
        ),
        (
            "/api/predict-risk",
            "Receives currency, amount, and horizon and returns risk, likelihood scores, cost estimates, drivers, and recommendations.",
        ),
        (
            "/api/feedback",
            "Provides the local participant-feedback backup endpoint.",
        ),
    ]
    for row, values in zip(tables[5].rows, function_rows):
        set_table_row(row, values)

    result_rows = [
        (
            "Currency",
            "Horizon",
            "Selected model",
            "Backtest accuracy",
            "Balanced accuracy",
            "Macro F1",
            "Holdout accuracy",
        ),
        ("USD/RWF", "7-day", "Logistic Regression", "0.5473", "0.5173", "0.4474", "1.0000*"),
        ("USD/RWF", "14-day", "Logistic Regression", "0.5963", "0.4420", "0.4044", "1.0000*"),
        ("EUR/RWF", "7-day", "Random Forest", "0.4136", "0.3462", "0.2667", "0.4691"),
        ("EUR/RWF", "14-day", "Logistic Regression", "0.3333", "0.3323", "0.2376", "0.6285"),
        ("KES/RWF", "7-day", "Logistic Regression", "0.4177", "0.3555", "0.2283", "0.4815"),
        ("KES/RWF", "14-day", "Logistic Regression", "0.4493", "0.3815", "0.2468", "0.4396"),
    ]
    for row, values in zip(tables[6].rows, result_rows):
        set_table_row(row, values)

    # Put the results table after the backtesting-method explanation.
    paragraphs[24]._p.addnext(tables[6]._tbl)
    insert_paragraphs_after(
        document,
        tables[6]._tbl,
        [
            (
                "Normal",
                (
                    "*The USD final holdouts contained only Low-risk observations: 324 "
                    "rows for 7 days and 323 rows for 14 days. Their 100% accuracy shows "
                    "correct classification of that stable period, but it does not prove "
                    "that the models distinguish all three risk classes."
                ),
            ),
            ("Heading 2", "Automated Software Testing"),
            (
                "Normal",
                (
                    "All 26 automated software tests passed. They covered official BNR "
                    "data validation, multicurrency API responses and predictions, Excel "
                    "report generation, model output labels, forward-only backtest folds, "
                    "purge gaps, final-holdout isolation, frontend structure, and account "
                    "security foundations. Passing these tests shows that the tested "
                    "software behaves as designed; it does not replace security auditing, "
                    "legal review, production monitoring, or future model validation."
                ),
            ),
            ("Heading 2", "Evaluation Conclusion"),
            (
                "Normal",
                (
                    "The testing provides a more realistic assessment than the earlier "
                    "single USD holdout. USD/RWF showed moderate historical signal, while "
                    "EUR/RWF and KES/RWF remained close to the three-class reference level. "
                    "FXGuard AI should therefore be described as a functioning and "
                    "transparent experimental decision-support prototype, not as a "
                    "guaranteed financial forecasting tool."
                ),
            ),
        ],
    )

    assumption_table = tables[8]
    assumption_updates = {
        2: (
            "Assumption 2",
            "USD, EUR, and KES represent useful supplier-payment currencies for the current prototype.",
        ),
        4: (
            "Limitation 1",
            "The current prototype supports USD, EUR, and KES only, rather than every currency used by Rwanda-based importers.",
        ),
        5: (
            "Limitation 2",
            "Backtest performance is mixed, and the final USD holdout contains only the Low class.",
        ),
        6: (
            "Limitation 3",
            "The tool provides experimental decision support, not guaranteed financial, forex-trading, or professional advice.",
        ),
        7: (
            "Future work",
            "Add macroeconomic features, live outcome monitoring, probability calibration, more currencies, scheduled data refresh, and repeated future validation.",
        ),
    }
    for row_index, (label, description) in assumption_updates.items():
        row = assumption_table.rows[row_index]
        replace_cell_text(row.cells[0], label)
        replace_cell_text(row.cells[1], label)
        replace_cell_text(row.cells[2], description)
        replace_cell_text(row.cells[3], description)

    document.core_properties.subject = (
        "Multicurrency exchange-rate risk forecasting, rolling-origin backtesting, "
        "and decision support"
    )
    document.core_properties.comments = (
        "Updated with verified multicurrency backtesting and automated test results."
    )

    document.save(TEMP_PATH)
    Document(TEMP_PATH)  # Reopen before replacement to verify the DOCX package.
    TEMP_PATH.replace(REPORT_PATH)
    print(f"Updated {REPORT_PATH.name}")
    print(f"Backup {BACKUP_PATH.name}")


if __name__ == "__main__":
    update_report()
